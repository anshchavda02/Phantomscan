"""
Builds an exceptionally formatted, publication-grade Microsoft Word (.docx) file
from PHANTOMSCAN_COMPLETE_TECHNICAL_REPORT.md.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Tuple, Optional

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Color Palette ---
COLOR_PRIMARY = RGBColor(0x0F, 0x17, 0x2A)     # Deep Navy
COLOR_SECONDARY = RGBColor(0x1E, 0x29, 0x3B)   # Slate Blue
COLOR_ACCENT = RGBColor(0x02, 0x84, 0xC7)      # Cyan / Ocean Blue
COLOR_TEXT = RGBColor(0x1E, 0x29, 0x3B)        # Charcoal Body Text
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)       # Muted Gray
COLOR_CODE = RGBColor(0x0F, 0x17, 0x2A)        # Code Text
COLOR_BORDER = "CBD5E1"                         # Light Gray Border
HEX_PRIMARY = "0F172A"
HEX_HEADER_BG = "1E293B"
HEX_ZEBRA = "F8FAFC"
HEX_CALLOUT_BG = "F1F5F9"
HEX_CALLOUT_BORDER = "0284C7"
HEX_CODE_BG = "F8FAFC"


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner margins (padding) for a table cell in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_background(cell, fill_hex: str):
    """Set background shading color of a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_table_borders(table, border_hex="CBD5E1"):
    """Apply clean, subtle borders to a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="{border_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{border_hex}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_hex}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def format_inline_runs(paragraph, text: str, font_name="Calibri", font_size=10.5, default_color=COLOR_TEXT, is_bold=False, is_italic=False):
    """Parse Markdown inline formatting (*italic*, **bold**, `code`, [link](url)) into formatted text runs."""
    # Pattern to match inline tokens
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
    
    for token in tokens:
        if not token:
            continue
        
        run = paragraph.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = default_color
        run.bold = is_bold
        run.italic = is_italic
        
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            run.text = token[2:-2]
            run.bold = True
            run.font.color.rgb = COLOR_PRIMARY
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            run.text = token[1:-1]
            run.italic = True
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(font_size - 0.5)
            run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
            # Add light shading to inline code
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
            rPr.append(shd)
        elif token.startswith('[') and '](' in token and token.endswith(')'):
            m = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if m:
                run.text = m.group(1)
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                run.underline = True
            else:
                run.text = token
        else:
            run.text = token


def add_code_block(doc, code_lines: list[str], language: str = ""):
    """Add a beautifully formatted code/ASCII box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    
    # Padding and styling
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_background(cell, HEX_CODE_BG)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{HEX_CALLOUT_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    
    full_text = "\n".join(code_lines)
    run = p.add_run(full_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    # Add a spacer paragraph after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)


def add_markdown_table(doc, table_lines: list[str]):
    """Convert Markdown table syntax into a styled Word table."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line or not line.startswith('|') or not line.endswith('|'):
            continue
        # Split by pipe, strip outer pipes
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Ignore separator row (| :--- | :---: |)
        if all(re.match(r'^:?-+:?$', c) for c in cells if c):
            continue
        rows.append(cells)
    
    if not rows:
        return
    
    num_cols = max(len(r) for r in rows)
    num_rows = len(rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "CBD5E1")
    
    # Repeat header row on every page
    header_tr = table.rows[0]._tr.get_or_add_trPr()
    header_tr.append(OxmlElement('w:tblHeader'))
    
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        
        # Don't split rows across pages
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        is_header = (r_idx == 0)
        
        for c_idx in range(num_cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
            
            if is_header:
                set_cell_background(cell, HEX_HEADER_BG)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                format_inline_runs(p, cell_text, font_name="Calibri", font_size=9.5, default_color=RGBColor(0xFF, 0xFF, 0xFF), is_bold=True)
            else:
                bg = HEX_ZEBRA if (r_idx % 2 == 1) else "FFFFFF"
                set_cell_background(cell, bg)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                format_inline_runs(p, cell_text, font_name="Calibri", font_size=9.5, default_color=COLOR_TEXT)

    # Spacer after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)


def convert_markdown_to_docx(md_path: Path, output_docx_path: Path):
    """Main converter reading Markdown and producing styled .docx."""
    doc = Document()
    
    # Configure 0.8 inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Configure Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("PhantomScan Complete Technical Feature & System Report  |  Confidential & Authorised Use Only")
        hrun.font.name = "Calibri"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun = fp.add_run("PhantomScan Cybersecurity Platform  •  https://github.com/anshchavda02/Phantomscan")
        frun.font.name = "Calibri"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = COLOR_MUTED

    # Configure Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = COLOR_TEXT
    normal_style.paragraph_format.line_spacing = 1.18
    normal_style.paragraph_format.space_after = Pt(4)

    lines = md_path.read_text(encoding='utf-8').splitlines()
    
    in_code_block = False
    code_block_lang = ""
    code_block_lines = []
    
    in_table = False
    table_lines = []
    
    for line in lines:
        stripped = line.strip()
        
        # --- Handle Code Block State ---
        if stripped.startswith('```'):
            if in_code_block:
                # Ending code block
                add_code_block(doc, code_block_lines, code_block_lang)
                in_code_block = False
                code_block_lines = []
                code_block_lang = ""
            else:
                # Starting code block
                if in_table:
                    add_markdown_table(doc, table_lines)
                    in_table = False
                    table_lines = []
                in_code_block = True
                code_block_lang = stripped[3:].strip()
                code_block_lines = []
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            continue
        
        # --- Handle Table State ---
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue
        else:
            if in_table:
                add_markdown_table(doc, table_lines)
                in_table = False
                table_lines = []
        
        # Blank line
        if not stripped:
            continue
        
        # Horizontal Rule (---)
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            continue
        
        # --- Headings ---
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            title_text = stripped[2:].strip()
            format_inline_runs(p, title_text, font_name="Calibri", font_size=18, default_color=COLOR_PRIMARY, is_bold=True)
            
            # Bottom border under main titles
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="{HEX_HEADER_BG}"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            continue

        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            h2_text = stripped[3:].strip()
            format_inline_runs(p, h2_text, font_name="Calibri", font_size=14, default_color=COLOR_PRIMARY, is_bold=True)
            continue

        if stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            h3_text = stripped[4:].strip()
            format_inline_runs(p, h3_text, font_name="Calibri", font_size=12, default_color=COLOR_ACCENT, is_bold=True)
            continue

        if stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            h4_text = stripped[5:].strip()
            format_inline_runs(p, h4_text, font_name="Calibri", font_size=11, default_color=COLOR_SECONDARY, is_bold=True)
            continue

        # --- Blockquotes / Alerts ---
        if stripped.startswith('>'):
            callout_text = stripped.lstrip('>').strip()
            t = doc.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = t.cell(0, 0)
            set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
            set_cell_background(cell, HEX_CALLOUT_BG)
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{HEX_CALLOUT_BORDER}"/>'
                f'  <w:top w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(borders)
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.line_spacing = 1.15
            format_inline_runs(cp, callout_text, font_name="Calibri", font_size=10, default_color=COLOR_SECONDARY, is_italic=True)
            
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(4)
            continue

        # --- Bullet Lists ---
        if re.match(r'^[-*]\s+', stripped):
            item_text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_inline_runs(p, item_text, font_name="Calibri", font_size=10.5, default_color=COLOR_TEXT)
            continue

        # --- Numbered Lists ---
        if re.match(r'^\d+\.\s+', stripped):
            item_text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_inline_runs(p, item_text, font_name="Calibri", font_size=10.5, default_color=COLOR_TEXT)
            continue

        # --- Standard Paragraph ---
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        format_inline_runs(p, stripped, font_name="Calibri", font_size=10.5, default_color=COLOR_TEXT)

    # Flush remaining table or code block
    if in_table:
        add_markdown_table(doc, table_lines)
    if in_code_block:
        add_code_block(doc, code_block_lines, code_block_lang)

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))
    print(f"[+] Successfully generated: {output_docx_path}")


if __name__ == '__main__':
    root = Path(__file__).resolve().parent.parent
    md_file = root / "PHANTOMSCAN_COMPLETE_TECHNICAL_REPORT.md"
    docx_file = root / "PHANTOMSCAN_COMPLETE_TECHNICAL_REPORT.docx"
    convert_markdown_to_docx(md_file, docx_file)
